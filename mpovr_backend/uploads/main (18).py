# Standard library imports
import os
import random
import logging
import json
import csv
import io
import zipfile
import shutil
from datetime import datetime, date, timedelta
from functools import wraps
from concurrent.futures import ProcessPoolExecutor
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import smtplib

# Third-party library imports
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_from_directory, send_file, Response,make_response
import requests
import pandas as pd
import pytz
import pyotp
import schedule
from io import BytesIO
from dateutil import parser
import time
from fos_chatbot import ask_question
import atexit
import threading
from werkzeug.serving import is_running_from_reloader
import xlsxwriter

# Your code using BytesIO here
# Local imports
from database import (
    get_data_from_database_for_download, verify_user, insert_log_to_db, get_buyer_name,
    store_in_process_data, bounce_complaint_action, get_company_id_by_fos_id,
    get_data_from_database_for_cs, get_all_logs, generate_and_save_qr_image,
    get_specific_complaint_data, get_app_token, fetch_report_data, generate_report_csv,
    get_all_surveys_for_dashboard, get_employee_stats_for_survey, fetch_all_departments,
    get_all_survey_questions, get_all_surveys_for_crud, add_question_to_db,
    update_question_in_db, delete_question_from_db, fetch_surveys, fetch_survey,
    fetch_survey_questions, fetch_question_responses, create_chart, fetch_survey_data,
    toggle_complaint_status, get_data_from_database_for_auditor, fetch_all_offices,
    edit_survey_in_db, delete_survey_from_db, add_new_survey, get_all_surveys,
    insert_complaint, generate_ticket, retrieve_employee_data,
    get_employee_data_for_fos_card, generate_anonymous_ticket, fetch_company_names,
    get_data_from_database, close_complaint, get_data_from_database_for_buyer,
    approve_complaint, reject_complaint, update_office_data, get_data_for_dashboard,
    add_new_office, get_all_company_names, update_notifications, get_office_id_from_ticket,
    get_app_token_from_ticket, delete_office_from_database, get_all_buyers_data,
    add_buyer_company, update_buyer_name, add_new_buyer, delete_buyers, get_office_name,
    get_department_names1, get_data_of_unregistered_complaints, insert_complaints,
    update_employee_in_database, is_employee_present, update_employee_left_in_database,
    get_last_fos_id, generate_leavers_joiners, update_cnic_numbers,
    insert_employees_to_database, insert_complaints_to_database, update_employees_data,
    add_notifications_for_late_capas, get_io_phone_no, retrieve_all_employee_data,
    get_notifications, add_notifications, update_employee_data, generate_employee_ids,
    add_employee_to_database, delete_employee, get_all_company_data, update_company_data,
    delete_company_data, add_new_company, get_all_offices_data, fetch_all_cache_employee_data,
    get_io_wise_complaint_count,check_and_send_reminders,fetch_all_buyers,fetch_companies_for_buyer,
    check_judge_grading_status,save_judge_responses,fetch_judge_responses,check_judge_consent,
    update_judge_consent,get_or_generate_summary
)
from sms_service import send_sms, send_otp, send_sms_to_io, send_login_otp, send_security_alert_email,send_rca_capa_reminder_email,send_factory_form_email,send_contact_form_email
from report_file import modify_survey_image, add_survey_details
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx import Document
from docx.shared import Inches, Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.events import EVENT_JOB_EXECUTED, EVENT_JOB_ERROR


from jj_api import insert_employees_from_api
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize,sent_tokenize
from nltk.probability import FreqDist
from heapq import nlargest
import spacy
from flask_cors import CORS  # Import CORS

try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    print("Downloading required NLTK data...")
    nltk.download('punkt')
    nltk.download('stopwords')
nltk.download('punkt_tab')
today = date.today()
logging.basicConfig(filename='logging.log', level=logging.DEBUG,format='%(asctime)s %(levelname)s %(message)s')

app = Flask(__name__)
app.config['SECRET_KEY'] = 'a599795a87199af397cb1e1e0c977687'
app.config['PERMANENT_SESSION_LIFETIME'] = 7200
app.secret_key = 'a599795a87199af397cb1e1e0c977687'
CORS(app, resources={
    r"/*": {
        "origins": [
            "https://nutribizpk.com",  # Specific origin you mentioned
            "http://localhost:3000",   # Common local development origin
            "https://localhost:3000",  # Secure local development origin
            "*"  # Be cautious with this in production
        ]
    }
})
pakistani_tz = pytz.timezone('Asia/Karachi')
validJudges = [
  "AbdulRehman", "SarimMehmood", "MiqdamJunaid", "SamanAslam", 
  "SarahBlanchard", "KarlBorgschulze", "RittaShine", "SagarMehmood", 
  "BariraHanif", "HaziqAhmed", "SamanHaseeb", "BadarUzaman", 
  "NaeemQureshi", "ShahbazSharif","JustajuVentures"
];
def send_otp_to_email(otp):
    try:
        message = f"FOS-Portal login: {otp}"
        send_otp('923120614727',message)
    except Exception as e:
        print("Error sending OTP email:", e)


def personal_access_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'email' in session and session['role'] == 'personal' and 'otp_verified' in session and session['otp_verified']:
            return f(*args, **kwargs)
        return redirect(url_for('login'))
    return decorated_function
def cs_access_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'email' in session and session['role'] == 'cs':
            return f(*args, **kwargs)
        return redirect(url_for('login'))
    return decorated_function
# Define a decorator to check if the user is authenticated and authorized to access the admin portal
def admin_access_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'email' in session and session['role'] == 'admin':
            return f(*args, **kwargs)
        return redirect(url_for('login'))
    return decorated_function

# Define a decorator to check if the user is authenticated and authorized to access the io portal
def io_access_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'email' in session and session['role'] == 'io':
            return f(*args, **kwargs)
        return redirect(url_for('login'))
    return decorated_function
def auditor_access_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'email' in session and session['role'] == 'auditor':
            return f(*args, **kwargs)
        return redirect(url_for('login'))
    return decorated_function
def handle_session_error(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        try:
            return f(*args, **kwargs)
        except KeyError as e:
            if str(e) in ["'email'", "'access_id'"]:
                # Log the error if needed
                print(f"Session error: {str(e)}")
                # Clear the session
                session.clear()
                # Redirect to login page
                return redirect(url_for('login'))
            else:
                # Re-raise the exception if it's not related to 'email' or 'access_id'
                raise
    return decorated_function

@app.route("/grading/<judge_name>", methods=['POST'])
def save_grading_responses(judge_name: str):
    """
    Route to save judge's grading responses for a participant.
    
    Args:
        judge_name (str): Name of the judge from URL path
    
    Returns:
        flask.Response: JSON response indicating success or failure
    """
    # Validate judge name
    if judge_name not in validJudges:
        return jsonify({"message": "Invalid judge name."}), 400
    
    # Parse request data
    data = request.get_json()
    
    # Validate participant name
    participant_name = data.get('participantName')
    if not participant_name or not isinstance(participant_name, str):
        return jsonify({"message": "Invalid participant name."}), 400
    
    # Validate responses
    responses = data.get('responses')
    if not responses or not isinstance(responses, list) or len(responses) != 5:
        return jsonify({
            "message": "Invalid responses format. Ensure you have exactly 5 responses."
        }), 400
    
    # Check if judge has already graded this participant
    if check_judge_grading_status(judge_name, participant_name):
        return jsonify({"message": "You have already graded this participant."}), 400
    
    # Save responses
    if save_judge_responses(judge_name, participant_name, responses):
        return jsonify({"message": "Responses saved successfully."}), 200
    else:
        return jsonify({"message": "Error saving responses."}), 500

@app.route("/grading/<judge_name>/responses", methods=['GET'])
def get_judge_responses(judge_name: str):
    """
    Route to retrieve all responses for a specific judge.
    
    Args:
        judge_name (str): Name of the judge from URL path
    
    Returns:
        flask.Response: JSON response with judge's responses or error message
    """
    # Check if the judgeName is valid
    if judge_name not in validJudges:
        return jsonify({"message": "Invalid judge name."}), 400
    
    # Fetch responses from the database
    responses = fetch_judge_responses(judge_name)
    
    # Handle potential database errors
    if responses is None:
        return jsonify({"message": "Error fetching responses."}), 500
    
    # Return responses
    return jsonify({"responses": responses}), 200
@app.route('/judges/consent/<judge_name>', methods=['GET'])
def get_judge_consent(judge_name: str):
    """
    Route to retrieve consent status for a specific judge.
    
    Args:
        judge_name (str): Name of the judge from URL path
    
    Returns:
        flask.Response: JSON response with consent status or error message
    """
    # Check consent status
    consent_status = check_judge_consent(judge_name)
    
    # Handle potential database errors
    if consent_status is None:
        return jsonify({"message": "Error checking consent status."}), 500
    
    # Return consent status
    return jsonify({"consentGiven": consent_status}), 200
@app.route('/judges/consent', methods=['POST'])
def record_judge_consent():
    """
    Route to record or update judge's consent status.
    
    Returns:
        flask.Response: JSON response indicating success or error
    """
    # Get request data
    data = request.get_json()
    
    # Extract judge name and consent status
    judge_name = data.get('judgeName')
    consent_given = data.get('consentGiven')
    
    # Validate judge name
    if judge_name not in validJudges:
        return jsonify({"message": "Invalid judge name."}), 400
    
    # Validate consent status
    if not isinstance(consent_given, bool):
        return jsonify({"message": "Invalid consent status."}), 400
    
    # Update consent status
    result = update_judge_consent(judge_name, consent_given)
    
    # Handle potential errors
    if result is None:
        return jsonify({
            "message": "Error recording consent.",
            "error": "Database operation failed"
        }), 500
    
    # Return success response
    return jsonify({"message": "Consent recorded successfully."}), 200
@app.route('/', methods=['GET', 'POST'])
def login():
    try:
        if request.method == 'POST':
            email = request.form['username']
            password = request.form['password']

            # Get user IP address
            user_ip = request.remote_addr
            
            # Get device name (user-agent)
            device_name = request.user_agent.platform
            print('Device Name:',device_name)
            user_info = verify_user(email, password)
            print('User Info:',user_info)
            if user_info:
                session['email'] = email
                session['role'] = user_info[0]
                session['access_id'] = user_info[1]
                print('User Info',user_info)
                

                # Log successful login attempt along with user IP address and device name
                app.logger.info(f"Successful login for email: {email}, role: {user_info[0]}, IP: {user_ip}, Device: {device_name}")

                # Check if the user is an admin or io, then redirect to respective portal
                if user_info[0] == 'admin':
                    if user_info[1] == 128:
                        otp = generate_otp()
                        send_login_otp(otp,number='923120614727')
                        session['otp'] = otp
                        session['otp_verified'] = False

                        return render_template('otp_verification.html', email=email)
                    return redirect(url_for('admin_portal'))
                elif user_info[0] == 'io':
                    return redirect(url_for('io_portal'))
                elif user_info[0] == 'cs':
                    return redirect(url_for('cs_table'))
                elif user_info[0] == 'auditor':
                    return redirect(url_for('survey_module'))
                elif user_info[0] == 'personal':
                    otp = generate_otp()
                    send_login_otp(otp)
                    session['otp'] = otp
                    session['email'] = email
                    session['role'] = 'personal'
                    session['otp_verified'] = False

                    # Log OTP generation and sending along with user IP address and device name
                    app.logger.info(f"OTP sent for email: {email}, IP: {user_ip}, Device: {device_name}")

                    return render_template('otp_verification.html', email=email)
                # Add more conditions for other roles if needed

                return redirect(url_for('success'))
            else:
                # Log failed login attempt along with user IP address and device name
                app.logger.warning(f"Failed login attempt for email: {email}, IP: {user_ip}, Device: {device_name}")
                return redirect(url_for('login'))

        return render_template('login.html')

    except Exception as e:
        # Log any exceptions that occur during the process along with user IP address and device name
        app.logger.error(f"An error occurred during login: {str(e)}, IP: {user_ip}, Device: {device_name}")
        return redirect(url_for('login'))

@app.route('/verify_otp', methods=['GET', 'POST'])
def verify_otp():

    email = request.args.get('email')
    
    expected_otp = session.get('otp')
    access_id = session['access_id']
    if request.method == 'POST':
        entered_otp = request.form['otp']
        if entered_otp == expected_otp:
            if access_id == 128:
                session['email'] = email
                session['role'] = 'admin'
                session['otp_verified'] = True
                return redirect(url_for('admin_portal'))
            else:
                session['email'] = email
                session['role'] = 'personal'
                session['otp_verified'] = True
                #send_security_alert_email()

                return redirect(url_for('personal_dashboard'))
        else:
            return render_template('otp_verification.html', error=True)

    return render_template('otp_verification.html', email=email)


def generate_otp():
    return str(random.randint(100000, 999999))

@app.route('/get_user_notifications/<int:user_id>', methods=['GET'])
def get_user_notifications(user_id):
    result = get_notifications(user_id)
    if result:
        notifications = [{'id': row[0], 'message': row[1], 'user_id': row[2], 'created_at': row[3],'is_read':row[4]} for row in result]
        return jsonify(notifications)
    else:
        return jsonify([])
@app.route('/update_user_notifications', methods=['POST'])
def update_user_notifications():
    try:
        data = request.get_json()
        user_id = data.get('user_id')
        if update_notifications(user_id):
            return jsonify({'message': 'Notifications updated successfully'})
        else:
            return jsonify({'message': 'Error occurred while updating notifications'})
    except Exception as e:
        # Handle any exceptions
        print(f"An error occurred: {e}")
        return jsonify({'message': 'Error occurred while updating notifications'})

@app.route('/update_dashboard_data', methods=['GET'])
@admin_access_required
def update_dashboard_data():
    access_id = session['access_id']
    start_date = request.args.get('startDate')
    end_date = request.args.get('endDate')
    
    # Convert strings to datetime objects
    start_date = datetime.strptime(start_date, '%Y-%m-%d')
    end_date = datetime.strptime(end_date, '%Y-%m-%d')
    
    # If you need to ensure they are timezone naive
    pakistani_tz = pytz.timezone('Asia/Karachi')
    start_date = pakistani_tz.localize(start_date)
    end_date = pakistani_tz.localize(end_date)
    
    result_data,allTimeComplaints = get_data_for_dashboard(access_id, start_date, end_date)
    images = load_companies_images(result_data)
    print('Dashboard Data: ',result_data)
    return jsonify({'companies': result_data, 'images': images})
@app.route('/admin_portal')
@admin_access_required
def admin_portal():
    user_ip = request.remote_addr
    username = session['email']
    # Your admin portal logic here
    access_id = session['access_id']
    #insert_complaints()
    #insert_complaints_to_database()
    today_date = datetime.now(pakistani_tz)
    date_30_days_ago = today_date - timedelta(days=30)
    #companies = get_department_names1(access_id)\
    companies,allTimeComplaints = get_data_for_dashboard(access_id,date_30_days_ago,today_date)
    buyer_name = get_buyer_name(access_id)
    company_logo = get_company_logo(buyer_name)
    insert_log_to_db(str(access_id),'Admin Dashboard','accessed',user_ip,username,'Logged In',buyer_name)
    buyer_name = buyer_name.split(" ", 2)[0]
    data = get_data_from_database_for_buyer(access_id)
    surveys = get_all_surveys_for_dashboard(access_id)
    #data = load_complaints_images(data)
    images = load_companies_images(companies)
    print
    if access_id in [127,128]:
            return render_template('sadaqat_dashboard.html', companies=companies,data=data, buyer_name=buyer_name,images=images,company_logo=company_logo,access_id=access_id,surveys=surveys,allTimeComplaints=allTimeComplaints)
    if access_id == 166:
            return render_template('lahore_dashboard.html', companies=companies,data=data, buyer_name=buyer_name,images=images,company_logo=company_logo,access_id=access_id,surveys=surveys,allTimeComplaints=allTimeComplaints)
    if access_id == 182:
            return render_template('fos_dashboard.html', companies=companies,data=data, buyer_name=buyer_name,images=images,company_logo=company_logo,access_id=access_id,surveys=surveys,allTimeComplaints=allTimeComplaints)
    if access_id in [15,16]:
        dormitorySummary = getDormitoryComplaintsSummary(data)
        return render_template('Admin_portal.html', companies=companies,data=data, buyer_name=buyer_name,images=images,company_logo=company_logo,access_id=access_id,surveys=surveys,allTimeComplaints=allTimeComplaints,dormitory_summary=dormitorySummary)
    return render_template('Admin_portal.html', companies=companies,data=data, buyer_name=buyer_name,images=images,company_logo=company_logo,access_id=access_id,surveys=surveys,allTimeComplaints=allTimeComplaints)
@app.route('/fetch_surveys_for_app')
def fetch_surveys_for_app():
    buyer_id =  request.args.get('buyerId')
    buyer_id = int(buyer_id)
    if buyer_id == 15:
        buyer_id = 16
    surveys = get_all_surveys_for_dashboard(buyer_id)
    return jsonify(surveys)
@app.route('/get_feedback_summary')
@admin_access_required
def get_feedback_summary():
    access_id = session['access_id']
    data = get_data_from_database_for_buyer(access_id)
    feedbackSummary = getFeedbackComplaintsSummary(data)
    return jsonify(feedbackSummary)

@app.route('/get_io_complaints', methods=['GET'])
@admin_access_required
def get_io_complaints():
    access_id = session['access_id']
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    start_date = parser.parse(start_date)
    end_date = parser.parse(end_date)
    data = get_io_wise_complaint_count(access_id, start_date, end_date)
    return jsonify(data)
@app.route('/download_reference_data')
def download_reference_data():
    access_id = session['access_id']
    user_ip = request.remote_addr
    username = session['email']
    insert_log_to_db(str(access_id),'Dashboard Reference Data','requested',user_ip,username,'','')
    
    start_date = request.args.get('startDate')
    end_date = request.args.get('endDate')
    
    # Convert string dates to datetime objects
    start_date = datetime.strptime(start_date, '%Y-%m-%d')
    end_date = datetime.strptime(end_date, '%Y-%m-%d')
    
    # Fetch all three types of complaints
    regular_complaints, feedback_complaints, dormitory_complaints = get_data_from_database_for_download(
        session['access_id'],
        start_date,
        end_date
    )
    
    # Create Excel writer object
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        # Create sheets based on data availability
        # Regular complaints sheet is always created
        pd.DataFrame(regular_complaints).to_excel(
            writer, 
            sheet_name='Regular Complaints',
            index=False
        )
        
        # Only create feedback sheet if there are feedback complaints
        if feedback_complaints:
            pd.DataFrame(feedback_complaints).to_excel(
                writer, 
                sheet_name='Feedback Complaints',
                index=False
            )
            
        # Only create dormitory sheet if there are dormitory complaints
        if dormitory_complaints:
            pd.DataFrame(dormitory_complaints).to_excel(
                writer, 
                sheet_name='Dormitory Complaints',
                index=False
            )
        
        # Auto-adjust column widths for each sheet
        for sheet_name in writer.sheets:
            worksheet = writer.sheets[sheet_name]
            if sheet_name == 'Regular Complaints':
                df = pd.DataFrame(regular_complaints)
            elif sheet_name == 'Feedback Complaints':
                df = pd.DataFrame(feedback_complaints)
            else:  # Dormitory Complaints
                df = pd.DataFrame(dormitory_complaints)
            
            for idx, col in enumerate(df.columns):
                max_length = max(
                    df[col].astype(str).apply(len).max(),
                    len(str(col))
                ) + 2
                worksheet.set_column(idx, idx, min(max_length, 50))  # Cap width at 50
    
    # Create response
    output.seek(0)
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'complaints_data_{start_date.date()}_{end_date.date()}.xlsx'
    )
@app.route('/get_buyers')
@admin_access_required
def get_buyers():
    buyers = fetch_all_buyers()
    if buyers is None:
        return jsonify({'error': 'Failed to fetch buyers'}), 500
    return jsonify(buyers)

@app.route('/get_companies/<int:buyer_id>')
@admin_access_required
def get_companies(buyer_id):
    companies = fetch_companies_for_buyer(buyer_id)
    if companies is None:
        return jsonify({'error': f'Failed to fetch companies for buyer {buyer_id}'}), 500
    return jsonify(companies)
@app.route('/survey_module')
@auditor_access_required
def survey_module():
    access_id = session['access_id']
    data = get_data_from_database_for_auditor(access_id)
    access_id = session['access_id']
    user_ip = request.remote_addr
    username = session['email']
    insert_log_to_db(str(access_id),'Handle Complaints Module','accessed',user_ip,username,'','')
    return render_template('handleComplaints.html', data=data)
@app.route('/toggle_complaint', methods=['POST'])
@auditor_access_required
def toggle_complaint():
    data = request.json
    ticket_number = data['ticket_number']
    action = data['action']
    success, message = toggle_complaint_status(ticket_number, action)
    access_id = session['access_id']
    user_ip = request.remote_addr
    username = session['email']
    insert_log_to_db(str(access_id),'Toggle Complaint Request',action,user_ip,username,'',ticket_number)
    return jsonify({"success": success, "message": message})
@app.route('/survey_dashboard')
@auditor_access_required
def survey_dashboard():
    access_id = session['access_id']
    survey_data = fetch_survey_data(access_id)
    user_ip = request.remote_addr
    username = session['email']
    insert_log_to_db(str(access_id),'Survey Dashboard','accessed',user_ip,username,'','')
    return render_template('surveyDashboard.html', survey_data=survey_data)
    
@app.route('/survey_report_dashboard', methods=['GET', 'POST'])
@auditor_access_required
def survey_report_dashboard():
    access_id = session['access_id']
    surveys = fetch_surveys(access_id)
    user_ip = request.remote_addr
    username = session['email']
    insert_log_to_db(str(access_id),'Survey Report Dashboard','accessed',user_ip,username,'','')
    return render_template('surveyReportDashboard.html', surveys=surveys)
@app.route('/survey/<int:survey_id>')
@auditor_access_required
def survey_report(survey_id):
    survey = fetch_survey(survey_id)
    questions = fetch_survey_questions(survey_id)
    charts = []
    text_responses = []

    for question in questions:
        responses = fetch_question_responses(survey_id, question['question_id'])
        if responses:
            if question['question_type'] in ['radio', 'select', 'checkbox']:
                fig = create_chart(question, responses)
                if fig:
                    chart_html = fig.to_html(full_html=False)
                    charts.append({'question': question, 'chart': chart_html})
            elif question['question_type'] in ['text', 'textarea']:
                text_responses.append({'question': question, 'responses': responses})
    access_id = session['access_id']
    user_ip = request.remote_addr
    username = session['email']
    insert_log_to_db(str(access_id),'Survey Report Request','Report',user_ip,username,'',str(survey_id))
    return render_template('survey_report.html', survey=survey, charts=charts, text_responses=text_responses)
@app.route('/survey/<int:survey_id>/report/csv')
def download_report_csv(survey_id):
    csv_data = generate_report_csv(survey_id)
    if not csv_data:
        return 'Survey data not found', 404

    # Create a response with the correct headers
    response = make_response(csv_data)
    response.headers['Content-Type'] = 'text/csv; charset=utf-8-sig'
    response.headers['Content-Disposition'] = f'attachment; filename=survey_{survey_id}_report.csv'

    return response
def fix_cs_formatting_runs(run_to_fix, user_cs_font_size, user_cs_font_name, user_is_bold):
    rpr = run_to_fix.element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rpr.get_or_add_sz()
    szCs = OxmlElement('w:szCs')  # size for complex script
    sz = OxmlElement('w:sz')  # size for other scripts
    rpr.append(szCs)
    rpr.append(sz)
    lang = OxmlElement('w:lang')  # language setting
    rpr.append(lang)
    if user_is_bold:
        bCs = OxmlElement('w:bCs')  # bold for complex script
        rpr.append(bCs)
        bCs.set(qn('w:val'), "True")
        b = OxmlElement('w:b')  # bold for other scripts
        rpr.append(b)
        b.set(qn('w:val'), "True")
    sz.set(qn('w:val'), str(int(user_cs_font_size * 2)))
    szCs.set(qn('w:val'), str(int(user_cs_font_size * 2)))
    lang.set(qn('w:bidi'), 'ar-SA')  # set language to Arabic
    rFonts.set(qn('w:cs'), user_cs_font_name)  # set complex script font
    rFonts.set(qn('w:ascii'), user_cs_font_name)  # set font for other scripts
    rFonts.set(qn('w:hAnsi'), user_cs_font_name)
def is_urdu(text):
    """Check if the text contains Urdu characters."""
    return any('\u0600' <= c <= '\u06FF' for c in text)
@app.route('/survey/<int:survey_id>/report/pdf')
def generate_pdf(survey_id):
    try:
        survey = fetch_survey(survey_id)
        questions = fetch_survey_questions(survey_id)
        image = get_company_logo_for_report(survey['buyer_name'])
        print('Report Image',image)
        employee_data = get_employee_stats_for_survey(survey_id)
        report_id = modify_survey_image(
            input_image_path="employee_feedback_survey.jpg",
            output_image_path="modified_survey_report.jpg",
            font_path="static/calibrib.ttf",
            logo_path=str(image),
            generated_for=survey['buyer_name'],
            buyer_id = survey['buyer_id']
        )
        
        input_image = "survey_details_image.jpg"
        output_image = "survey_details_filled.jpg"
        add_survey_details(input_image, output_image, survey, employee_data)

        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Jameel Noori Nastaleeq'
        font.size = Pt(12)

        # Set the page layout for the entire document
        sections = document.sections
        for section in sections:
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)

        # Add images to the document
        base_path="static/survey_reports/images"
        document.add_picture(os.path.join(base_path,"modified_survey_report.jpg"), width=Inches(8.27))
        document.add_section()
        document.add_picture(os.path.join(base_path,"survey_details_filled.jpg"), width=Inches(8.27))

        # Add a new section for adding headers and footers
        document.add_section()
        for i, section in enumerate(document.sections):
            if i >= 2:
                header = section.header
                footer = section.footer
                section.header_distance = Inches(0)
                section.footer_distance = Inches(0)
                
                # Add header and footer images
                header_paragraph = header.paragraphs[0]
                header_paragraph.add_run().add_picture('header.jpg', width=Inches(8.27))

                footer_paragraph = footer.paragraphs[0]
                footer_paragraph.add_run().add_picture('footer.jpg', width=Inches(8.27))
            else:
                section.different_first_page_header_footer = True
        heading = document.add_paragraph()
        heading_run = heading.add_run('Survey Results')
        heading_run.font.name = 'Arial'  # You can change this to any preferred font
        heading_run.font.size = Pt(24)  # Adjust size as needed
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0x2D, 0x94, 0x80)  # Color #2D9480
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.left_indent = Inches(0.5)

        # Add some space after the heading
        heading.paragraph_format.space_after = Pt(12)
        # Add a table for survey title and description
        table = document.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set the first row as the header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Title'
        header_cells[1].text = 'Description'
        header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set the survey title and description in the table
        data_cells = table.rows[1].cells
        data_cells[0].text = survey['title']
        data_cells[1].text = survey['description']

        for cell in table.columns[0].cells:
            cell.width = Inches(3)
        color = RGBColor(0x28, 0x49, 0x52)
        for cell in data_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = color
        # Apply RTL formatting if text is in Urdu
        for cell in data_cells:
            for paragraph in cell.paragraphs:
                if is_urdu(paragraph.text):
                    for run in paragraph.runs:
                        run.font.rtl = True  # Apply RTL
                        fix_cs_formatting_runs(run, 12, "Jameel Noori Nastaleeq", False)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add a new style for spacing
        spacing = document.styles.add_style('Spacing', WD_STYLE_TYPE.PARAGRAPH)
        spacing.paragraph_format.space_after = Pt(24)

        # Track the number of charts on the current page
        chart_count = 0
        first_chart_placed = False

        # Add a counter for question numbers
        question_number = 1

        for question in questions:
            question_text = f"Q{question_number}. {question['question_text']}"
            p = document.add_paragraph(question_text)
            r = p.runs[0]
            font = r.font
            font.bold = True
            font.color.rgb = RGBColor(0xf5, 0xa8, 0x3c)
            if is_urdu(question['question_text']):
                r.font.rtl = True  # Apply RTL
                fix_cs_formatting_runs(r, 20, "Jameel Noori Nastaleeq", False)  # Use appropriate Urdu font and settings
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.right_indent = Inches(0.5)
            else:
                font.complex_script = False
                font.rtl = False
                font.name = 'Calibri'
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Inches(0.5)
              # Adjust the value as needed

            responses = fetch_question_responses(survey_id, question['question_id'])

            if question['question_type'] in ['radio', 'select', 'checkbox']:
                img_buffer = create_chart(question, responses, f"Number of responses: {len(responses) }")
                if img_buffer:
                    # Add chart with center alignment
                    chart_paragraph = document.add_paragraph()
                    run = chart_paragraph.add_run()
                    run.add_picture(img_buffer, width=Inches(6))
                    chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if not first_chart_placed:
                        # If it's the first chart, place it on a single page
                        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                        first_chart_placed = True
                    else:
                        # Increment chart count
                        chart_count += 1

                        # Add a page break after every two charts
                        if chart_count == 2:
                            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                            chart_count = 0  # Reset chart count

            elif question['question_type'] in ['text', 'textarea']:
                # Add a subheading for responses
                response_heading = document.add_paragraph()
                response_run = response_heading.add_run("Responses:")
                response_run.font.name = 'Calibri'
                response_run.font.size = Pt(14)
                response_run.font.bold = True
                response_run.font.color.rgb = RGBColor(0x2D, 0x94, 0x80)  # Teal color
                response_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                response_heading.paragraph_format.left_indent = Inches(0.5)
                response_heading.paragraph_format.space_after = Pt(6)

                for index, response in enumerate(responses, 1):
                    p = document.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.7)
                    p.paragraph_format.right_indent = Inches(0.7)
                    p.paragraph_format.space_after = Pt(12)
                    
                    # Add anonymous response number
                    r = p.add_run(f"Response {index}: ")
                    r.font.name = 'Calibri'
                    r.font.size = Pt(11)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0x28, 0x49, 0x52)  # Dark teal

                    # Add response text
                    r = p.add_run(response['answer_text'])
                    r.font.name = 'Calibri' if not is_urdu(response['answer_text']) else 'Jameel Noori Nastaleeq'
                    r.font.size = Pt(11)
                    
                    if is_urdu(response['answer_text']):
                        r.font.complex_script = True
                        r.font.rtl = True
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        fix_cs_formatting_runs(r, 11, "Jameel Noori Nastaleeq", False)
                    else:
                        r.font.complex_script = False
                        r.font.rtl = False
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Add a light gray background to the paragraph
                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
                    p._element.get_or_add_pPr().append(shading_elm)

                # Add some space after the responses
                document.add_paragraph().paragraph_format.space_after = Pt(12)

            # Increment the question number counter
            question_number += 1

        file_buffer = BytesIO()
        document.save(file_buffer)
        file_buffer.seek(0)
        
        word_file_path = f'static/survey_reports/{report_id}.docx'
        document.save(word_file_path)
        if survey_id == 12:
            pdf_file_path = f'static/survey_reports/Survey12 report - Analysis English.pdf'  # Adjust this path as needed
    
            # Create a ZIP file in memory
            memory_file = BytesIO()
            with zipfile.ZipFile(memory_file, 'w') as zf:
                # Add Word document to the ZIP
                zf.write(word_file_path, f'survey_{survey_id}_report.docx')
                
                # Add PDF to the ZIP if it exists
                if os.path.exists(pdf_file_path):
                    zf.write(pdf_file_path, f'survey_{survey_id}_english.pdf')
                else:
                    print(f"PDF file not found: {pdf_file_path}")
    
            # Seek to the beginning of the BytesIO object
            memory_file.seek(0)
    
            # Send the ZIP file
            return send_file(
                memory_file,
                mimetype='application/zip',
                as_attachment=True,
                download_name=f'survey_{survey_id}_report.zip'
            )
        
        return send_file(file_buffer, as_attachment=True, download_name=f'survey_{survey_id}_report.docx')
    except Exception as e:
        print(f"Error generating Word report: {str(e)}")
        return 'Error generating report', 500

@app.route('/survey/<int:survey_id>/report/docx/app')
def generate_word_app(survey_id):
    try:
        survey = fetch_survey(survey_id)
        questions = fetch_survey_questions(survey_id)
        image = get_company_logo_for_report(survey['buyer_name'])
        print('Report Image',image)
        employee_data = get_employee_stats_for_survey(survey_id)
        report_id = modify_survey_image(
            input_image_path="employee_feedback_survey.jpg",
            output_image_path="modified_survey_report.jpg",
            font_path="static/calibrib.ttf",
            logo_path=str(image),
            generated_for=survey['buyer_name'],
            buyer_id = survey['buyer_id']
        )
        
        input_image = "survey_details_image.jpg"
        output_image = "survey_details_filled.jpg"
        add_survey_details(input_image, output_image, survey, employee_data)

        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Jameel Noori Nastaleeq'
        font.size = Pt(12)

        # Set the page layout for the entire document
        sections = document.sections
        for section in sections:
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)

        # Add images to the document
        base_path="static/survey_reports/images"
        document.add_picture(os.path.join(base_path,"modified_survey_report.jpg"), width=Inches(8.27))
        document.add_section()
        document.add_picture(os.path.join(base_path,"survey_details_filled.jpg"), width=Inches(8.27))

        # Add a new section for adding headers and footers
        document.add_section()
        for i, section in enumerate(document.sections):
            if i >= 2:
                header = section.header
                footer = section.footer
                section.header_distance = Inches(0)
                section.footer_distance = Inches(0)
                
                # Add header and footer images
                header_paragraph = header.paragraphs[0]
                header_paragraph.add_run().add_picture('header.jpg', width=Inches(8.27))

                footer_paragraph = footer.paragraphs[0]
                footer_paragraph.add_run().add_picture('footer.jpg', width=Inches(8.27))
            else:
                section.different_first_page_header_footer = True
        heading = document.add_paragraph()
        heading_run = heading.add_run('Survey Results')
        heading_run.font.name = 'Arial'  # You can change this to any preferred font
        heading_run.font.size = Pt(24)  # Adjust size as needed
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0x2D, 0x94, 0x80)  # Color #2D9480
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.left_indent = Inches(0.5)

        # Add some space after the heading
        heading.paragraph_format.space_after = Pt(12)
        # Add a table for survey title and description
        table = document.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set the first row as the header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Title'
        header_cells[1].text = 'Description'
        header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set the survey title and description in the table
        data_cells = table.rows[1].cells
        data_cells[0].text = survey['title']
        data_cells[1].text = survey['description']

        for cell in table.columns[0].cells:
            cell.width = Inches(3)
        color = RGBColor(0x28, 0x49, 0x52)
        for cell in data_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = color
        # Apply RTL formatting if text is in Urdu
        for cell in data_cells:
            for paragraph in cell.paragraphs:
                if is_urdu(paragraph.text):
                    for run in paragraph.runs:
                        run.font.rtl = True  # Apply RTL
                        fix_cs_formatting_runs(run, 12, "Jameel Noori Nastaleeq", False)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add a new style for spacing
        spacing = document.styles.add_style('Spacing', WD_STYLE_TYPE.PARAGRAPH)
        spacing.paragraph_format.space_after = Pt(24)

        # Track the number of charts on the current page
        chart_count = 0
        first_chart_placed = False

        # Add a counter for question numbers
        question_number = 1

        for question in questions:
            question_text = f"Q{question_number}. {question['question_text']}"
            p = document.add_paragraph(question_text)
            r = p.runs[0]
            font = r.font
            font.bold = True
            font.color.rgb = RGBColor(0xf5, 0xa8, 0x3c)
            if is_urdu(question['question_text']):
                r.font.rtl = True  # Apply RTL
                fix_cs_formatting_runs(r, 20, "Jameel Noori Nastaleeq", False)  # Use appropriate Urdu font and settings
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.right_indent = Inches(0.5)
            else:
                font.complex_script = False
                font.rtl = False
                font.name = 'Calibri'
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Inches(0.5)
              # Adjust the value as needed

            responses = fetch_question_responses(survey_id, question['question_id'])

            if question['question_type'] in ['radio', 'select', 'checkbox']:
                img_buffer = create_chart(question, responses, f"Number of responses: {len(responses) }")
                if img_buffer:
                    # Add chart with center alignment
                    chart_paragraph = document.add_paragraph()
                    run = chart_paragraph.add_run()
                    run.add_picture(img_buffer, width=Inches(6))
                    chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if not first_chart_placed:
                        # If it's the first chart, place it on a single page
                        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                        first_chart_placed = True
                    else:
                        # Increment chart count
                        chart_count += 1

                        # Add a page break after every two charts
                        if chart_count == 2:
                            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                            chart_count = 0  # Reset chart count

            elif question['question_type'] in ['text', 'textarea']:
                # Add a subheading for responses
                response_heading = document.add_paragraph()
                response_run = response_heading.add_run("Responses:")
                response_run.font.name = 'Calibri'
                response_run.font.size = Pt(14)
                response_run.font.bold = True
                response_run.font.color.rgb = RGBColor(0x2D, 0x94, 0x80)  # Teal color
                response_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                response_heading.paragraph_format.left_indent = Inches(0.5)
                response_heading.paragraph_format.space_after = Pt(6)

                for index, response in enumerate(responses, 1):
                    p = document.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.7)
                    p.paragraph_format.right_indent = Inches(0.7)
                    p.paragraph_format.space_after = Pt(12)
                    
                    # Add anonymous response number
                    r = p.add_run(f"Response {index}: ")
                    r.font.name = 'Calibri'
                    r.font.size = Pt(11)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0x28, 0x49, 0x52)  # Dark teal

                    # Add response text
                    r = p.add_run(response['answer_text'])
                    r.font.name = 'Calibri' if not is_urdu(response['answer_text']) else 'Jameel Noori Nastaleeq'
                    r.font.size = Pt(11)
                    
                    if is_urdu(response['answer_text']):
                        r.font.complex_script = True
                        r.font.rtl = True
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        fix_cs_formatting_runs(r, 11, "Jameel Noori Nastaleeq", False)
                    else:
                        r.font.complex_script = False
                        r.font.rtl = False
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Add a light gray background to the paragraph
                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
                    p._element.get_or_add_pPr().append(shading_elm)

                # Add some space after the responses
                document.add_paragraph().paragraph_format.space_after = Pt(12)

            # Increment the question number counter
            question_number += 1

        file_buffer = BytesIO()
        document.save(file_buffer)
        file_buffer.seek(0)
        
        word_file_path = f'static/survey_reports/{report_id}.docx'
        document.save(word_file_path)
        
        return send_file(file_buffer, as_attachment=True, download_name=f'survey_{survey_id}_report.docx')
    except Exception as e:
        print(f"Error generating Word report: {str(e)}")
        return 'Error generating report', 500
@app.route('/survey/<int:survey_id>/report/pdf/app')
def generate_pdf_app(survey_id):
    try:
        pdf_file_path = f'static/survey_reports/Survey12 report - Analysis English.pdf'  # Adjust this path as needed

        
        return send_file(pdf_file_path, as_attachment=True, download_name=f'survey_{survey_id}_report_analysis.pdf')
    except Exception as e:
        print(f"Error generating Word report: {str(e)}")
        return 'Error generating report', 500
@app.route('/io_portal')
@io_access_required
def io_portal():
    access_id = session['access_id']
    data = get_data_from_database(access_id)
    #data = load_complaints_images(data)
    office_name = get_office_name(access_id)
    if 'DHA' in office_name:
        office_name = 'Cheezious'
    elif 'Executive' in office_name or 'Export' in office_name or 'Human' in office_name or 'Finance' in office_name:
        office_name = 'Dawn'
    if access_id in range(212,221):
        office_name = 'Denim-E'
    if access_id == 149:
        office_name = 'Dormitory'
    user_ip = request.remote_addr
    username = session['email']
    insert_log_to_db(str(access_id),'IO Portal','accessed',user_ip,username,'Logged In',office_name)
    if office_name:
        office_name = office_name.split(' ',2)[0]
    if access_id == 211:
        return render_template('lahore_portal.html', data=data,office_name=office_name,access_id=access_id)

    return render_template('aas.html', data=data,office_name=office_name,access_id=access_id)
    
@app.route('/io_portal_json')
@io_access_required
def io_portal_json():
    access_id = session['access_id']
    data = get_data_from_database(access_id)
    return jsonify(data=data)
    
@app.route('/submit_form', methods=['POST'])
@io_access_required
def submit_form():
    try:
        ticket_number = request.form.get('ticket')
        rca = request.form.get('rca')
        capa = request.form.get('capa')
        rca1 = request.form.get('rca1')
        capa1 = request.form.get('capa1')
        rca2 = request.form.get('rca2')
        capa2 = request.form.get('capa2')
        unprocessed = request.form.get('unprocessed')
        rca_datetime = request.form.get('rcaDeadline')
        rca1_datetime = request.form.get('rca1Deadline')
        rca2_datetime = request.form.get('rca2Deadline')
        
        # Convert 'null' string to None
        fields = [ticket_number, rca, capa, rca1, capa1, rca2, capa2, unprocessed, rca_datetime, rca1_datetime, rca2_datetime]
        fields = [None if field == 'null' else field for field in fields]
        ticket_number, rca, capa, rca1, capa1, rca2, capa2, unprocessed, rca_datetime, rca1_datetime, rca2_datetime = fields
        username = session['email']
        user_ip = request.remote_addr
        access_id = session['access_id']
        # Process the data and store it as needed
        store_in_process_data(ticket_number, rca, capa, rca1, capa1, rca2, capa2, unprocessed,username,user_ip,access_id, rca_datetime, rca1_datetime, rca2_datetime)
        # Handle uploaded files
        capa_image = request.files.get('capaImage')
        capa1_image = request.files.get('capa1Image')
        capa2_image = request.files.get('capa2Image')
        # Save the uploaded images
        if capa_image:
            save_uploaded_image(ticket_number, capa_image, 'capa')
        if capa1_image:
            save_uploaded_image(ticket_number, capa1_image, 'capa1')
        if capa2_image:
            save_uploaded_image(ticket_number, capa2_image, 'capa2')

        # Return a success response as JSON
        return jsonify({'message': 'Form submitted successfully'}), 200

    except Exception as e:
        print('error', e)
        return jsonify({'error': str(e)}), 500

@app.route('/entry_form')
def cs_dashboard():
    company_names = fetch_company_names()
    return render_template('entry_form.html', company_names=company_names)

@app.route('/cs_entry')
@cs_access_required
def cs_entry():
    company_names = fetch_company_names()
    access_id = session['access_id']
    ip_address = request.remote_addr
    username = session['email']
    insert_log_to_db(str(access_id),'CS Entry Form','accessed',ip_address,username,'','')
    return render_template('cs_entry_form.html', company_names=company_names)
    
def getdate():
    today = datetime.now(pakistani_tz)
    day = f"0{today.day}"
    month = f"0{today.month}"
    curr_date = f"{day[-2:]}{month[-2:]}"
    return curr_date

checkbox_texts = {
    'option-1': 'Workplace Health, Safety and Environment',
    'option-2': 'Freedom of Association',
    'option-3': 'Child Labor',
    'option-4': 'Wages & Benefits',
    'option-5': 'Working Hours',
    'option-6': 'Forced Labor',
    'option-7': 'Discrimination',
    'option-8': 'Unfair Employment',
    'option-9': 'Ethical Business',
    'option-10': 'Harassment',
    'option-11': 'Workplace Discipline',
    'option-12': 'Feedback'
}


@app.route('/manage_complaints')
@personal_access_required
def manage_complaints():
    data = get_data_from_database_for_cs()
    data = load_complaints_images(data)
    return render_template('complaint_crud.html', data=data)
@app.route('/unregistered_complaints')
@personal_access_required
def unregistered_complaints():
    data = get_data_of_unregistered_complaints()
    return render_template('unregistered_complaints.html', data=data)
@app.route('/cs_table')
@cs_access_required
def cs_table():
    data = get_data_from_database_for_cs()
    #data = load_complaints_images(data)
    access_id = session['access_id']
    user_ip = request.remote_addr
    username = session['email']
    insert_log_to_db(str(access_id),'CS Portal','accessed',user_ip,username,'Logged In','')
    return render_template('cs_table.html', data=data)

@app.route('/cs_table_json')
@cs_access_required
def cs_table_json():
    data = get_data_from_database_for_cs()
    return jsonify(data=data)
@app.route('/submit', methods=['POST'])
def submit_form1():
    curr_time = datetime.now(pakistani_tz)
    curr_time = curr_time.strftime('%Y-%m-%d %H:%M:%S')
    user_ip = request.remote_addr

    if request.form.get('q14_company') == '':
        reference_number = 0
    else:
        reference_number = int(request.form.get('q14_company'))

    data = {
        'reference_number': reference_number,  # Assuming the reference number is an integer
        'is_urgent': 'button1' in request.form,
        'is_anonymous': 'button2' in request.form,
        'company_name': request.form.get('q_company'),
        'mobile_number': request.form.get('q_input4'),
        'date_of_issue': request.form.get('q_new_date'),
        'complaint_categories': request.form.getlist('complaint_category'),
        'additional_comments': request.form.get('q45_clickTo45'),
        'person_issue': request.form.get('q_person_issue'),
        'concerned_department': request.form.get('q_concerned_department'),
        'previous_history': request.form.get('q_previous_history'),
        'proposed_solution': request.form.get('q_proposed_solution'),
        'curr_date': curr_time,
        'status': 'Unapproved',
        'user_ip':user_ip

    }
    ticket = ""

    for complaints in data["complaint_categories"]:
        if data['is_anonymous'] == True and reference_number !=0:
            ticket_no = generate_anonymous_ticket(str(data['reference_number']), getdate())
        elif reference_number == 0:
            ticket_no = generate_anonymous_ticket('11',getdate())
        else:
            ticket_no = generate_ticket(checkbox_texts[complaints],getdate(),data['reference_number'])
        ticket += ticket_no + '<br>'
        data.update({'ticket_number': ticket_no})
        data.update({'complaint_categories':checkbox_texts[complaints]})
        
        if insert_complaint(data,0):
            print(data['ticket_number'])
            insert_log_to_db('','Entry Form','Complaint Registered',user_ip,'',data['additional_comments'].split('.', 1)[0],ticket_no)
            file = request.files['q_image_upload']
            if file:
                # Save the file to a designated folder
                save_uploaded_image(data['ticket_number'], file, 'proof')
                print('file Uploaded',file)
            
            #mobile_number,office_id = get_io_phone_no(data['reference_number'],data['complaint_categories'])
            #if data['is_anonymous'] == True or reference_number == 0:
            #    message = f"New Complaint #{data['ticket_number']} is registered from an anonymous employee."
            #else:
            #    message = f"New Complaint #{data['ticket_number']} is registered from employee #{reference_number}."

            #add_notifications(message,office_id)
            #send_sms_to_io(mobile_number, ticket_no)
            #send_notification(data['reference_number'],ticket_no)
            #send_sms(data['mobile_number'], ticket_no)
        else:
            return redirect(url_for('cs_dashboard'))
    return redirect(url_for('ticket_no', ticket_number=ticket))

@app.route('/cs_submit', methods=['POST'])
@cs_access_required
def cs_submit_form():
    user_ip = request.remote_addr
    access_id = session['access_id']
    username = session['email']
    curr_time = datetime.now(pakistani_tz)
    curr_time = curr_time.strftime('%Y-%m-%d %H:%M:%S')
    

    if request.form.get('q14_company') == '':
        reference_number = 0
    else:
        reference_number = int(request.form.get('q14_company'))

    data = {
        'reference_number': reference_number,  # Assuming the reference number is an integer
        'is_urgent': 'button1' in request.form,
        'is_anonymous': 'button2' in request.form,
        'company_name': request.form.get('q_company'),
        'mobile_number': request.form.get('q_input4'),
        'date_of_issue': request.form.get('q_new_date'),
        'complaint_categories': request.form.getlist('complaint_category'),
        'additional_comments': request.form.get('q45_clickTo45'),
        'person_issue': request.form.get('q_person_issue'),
        'concerned_department': request.form.get('q_concerned_department'),
        'previous_history': request.form.get('q_previous_history'),
        'proposed_solution': request.form.get('q_proposed_solution'),
        'curr_date': curr_time,
        'status': 'Unprocessed',
        'user_ip':user_ip

    }
    ticket = ""

    for complaints in data["complaint_categories"]:
        if data['is_anonymous'] == True and reference_number !=0:
            ticket_no = generate_anonymous_ticket(str(data['reference_number']), getdate())
        elif reference_number == 0:
            ticket_no = generate_anonymous_ticket('11',getdate())
        else:
            ticket_no = generate_ticket(checkbox_texts[complaints],getdate(),data['reference_number'])
        ticket += ticket_no + '<br>'
        data.update({'ticket_number': ticket_no})
        data.update({'complaint_categories':checkbox_texts[complaints]})
        
        if insert_complaint(data,1):
            insert_log_to_db('','CS Entry Form','Complaint Registered',user_ip,username,data['additional_comments'].split('.', 1)[0],ticket_no)
            file = request.files['q_image_upload']
            if file:
                # Save the file to a designated folder
                save_uploaded_image(data['ticket_number'], file, 'proof')
                print('file Uploaded',file)
            mobile_number,office_id = get_io_phone_no(data['reference_number'],data['complaint_categories'],data['additional_comments'])
            if data['is_anonymous'] == True or reference_number == 0:
                message = f"New Complaint #{data['ticket_number']} is registered from an anonymous employee."
            else:
                message = f"New Complaint #{data['ticket_number']} is registered from employee #{reference_number}."

            add_notifications(message,office_id)
            send_sms_to_io(mobile_number, ticket_no)
            send_notification(data['reference_number'],ticket_no)
            send_sms(data['mobile_number'], ticket_no)
        else:
            return redirect(url_for('cs_dashboard'))
    return redirect(url_for('ticket_no', ticket_number=ticket))
def send_notification( reference_number, ticket_number):
    EXPO_PUSH_ENDPOINT = 'https://exp.host/--/api/v2/push/send'
    app_token = get_app_token(reference_number)
    payload = {
        'to': app_token,
        'sound': 'default',
        'title': 'New Complaint Registered',
        'body': f'Complaint #{ticket_number} is registered from employee #{reference_number}',
        'data': {'ticket_number': ticket_number}
    }

    try:
        response = requests.post(EXPO_PUSH_ENDPOINT, json=payload)
        if response.status_code == 200:
            print('Notification sent successfully!')
        else:
            print('Error sending notification:', response.status_code)
    except Exception as e:
        print('Error sending notification:', str(e))
@app.route("/get_data/<string:reference_number>", methods=["GET"])
def get_data(reference_number):
    try:
        # Retrieve employee data, including additional fields
        employee_data = retrieve_employee_data(reference_number)
        if employee_data:
            # Create the response JSON data directly with the retrieved values
            data = {
                reference_number: {
                    "name": employee_data["employee_name"],
                    "company": employee_data["company_name"],
                    "workerType": employee_data["worker_type"],
                    "department": employee_data["department"],
                    "designation": employee_data["designation"],
                    "gender": employee_data["gender"],
                    "mobile": employee_data["mobile_number"],
                    "office_id": employee_data["office_id"],
                    "employee_left": employee_data["employee_left"],
                    "employeeId": employee_data["employee_id"]
                }
            }
            logging.info(f"Received request for reference number: {reference_number}, Data: {data}")
            return jsonify(data)
        else:
            return jsonify({})
    except Exception as e:
        logging.error(f"Error fetching data for reference number {reference_number}: {e}")
        return jsonify({})
        
@app.route('/thank_you/<ticket_number>')
def ticket_no(ticket_number):
    return render_template('ticket_no.html', ticket_number=ticket_number)

@app.route('/home')
def home():
    company_names = fetch_company_names()
    return render_template('entry_form.html', company_names=company_names)

@app.route('/manage_logs')
@personal_access_required
def manage_logs():
    try:
        logs = get_all_logs()
        return render_template('logs.html',logs = logs)
    except Exception as e:
        error_message = {'error': str(e)}
        return jsonify(error_message), 500

@app.route('/close_complaint', methods=['POST'])
@cs_access_required
def close_complaints():
    try:
        data = request.json
        user_id = session['access_id']
        username =  session['email']
        user_ip = request.remote_addr
        ticket_number = data.get('ticket')
        feedback = data.get('feedback')

        if close_complaint(ticket_number,feedback):
            insert_log_to_db('','CS Portal','Complaint Closed',user_ip,username,feedback.split('.', 1)[0],ticket_number)
            response_data = {'message': 'Complaint closed successfully'}
            return jsonify(response_data), 200
        else:
            response_data = {'message': 'Complaint could not be closed'}
            return jsonify(response_data), 400  # You can choose an appropriate status code here
    except Exception as e:
        # Handle any errors that might occur during processing
        error_message = {'error': str(e)}
        return jsonify(error_message), 500
@app.route('/bounce_complaint', methods=['POST'])
@cs_access_required
def bounce_complaint():
    try:
        data = request.json
        user_ip = request.remote_addr
        username = session['email']
        ticket_number = data.get('ticket')
        feedback = data.get('feedback')
        if bounce_complaint_action(ticket_number,feedback):
            insert_log_to_db('','CS Portal','Complaint Bounced',user_ip,username,feedback.split('.', 1)[0],ticket_number)
            office_id = get_office_id_from_ticket(ticket_number)
            response_data = {'message': 'Complaint bounced successfully'}
            print('Complaint bounced successfully')
            message = f"Complaint #{ticket_number} is bounced. The complainant is not satisfied."
            add_notifications(message,office_id)
            send_notification_for_bounced(ticket_number)
            return jsonify(response_data), 200
        else:
            response_data = {'message': 'Complaint bouncing failed'}
            return jsonify(response_data), 400  # You can choose an appropriate status code here
    except Exception as e:
        # Handle any errors that might occur during processing
        error_message = {'error': str(e)}
        return jsonify(error_message), 500

def send_notification_for_bounced( ticket_number):
    EXPO_PUSH_ENDPOINT = 'https://exp.host/--/api/v2/push/send'
    app_token = get_app_token_from_ticket(ticket_number)
    payload = {
        'to': app_token,
        'sound': 'default',
        'title': 'Complaint Bounced',
        'body': f'Complaint #{ticket_number} is bounced! Complainant not Satisfied.',
        'data': {'ticket_number': ticket_number}
    }

    try:
        response = requests.post(EXPO_PUSH_ENDPOINT, json=payload)
        if response.status_code == 200:
            print('Bounced Notification sent successfully!')
        else:
            print('Error sending Bounced notification:', response.status_code)
    except Exception as e:
        print('Error sending notification:', str(e))

def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'csv', 'xlsx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
@app.route('/personal_dashboard')
@personal_access_required
def personal_dashboard():
    data = get_all_company_data()
    user_ip = request.remote_addr
    username = session['email']

    insert_log_to_db('','Personal Account','Accessed',user_ip,username,'Logged In','')
    return render_template('company_crud.html', data=data)

@app.route('/manage_company')
@personal_access_required
def manage_company():
    data = get_all_company_data()
    return render_template('company_crud.html', data=data)



@app.route('/manage_employee')
@personal_access_required
def manage_employee():
    return render_template('employee_crud.html')
@app.route('/get_employees_data', methods=['POST'])
@personal_access_required
def get_employees_data():

    # Get the parameters from DataTables
    draw = int(request.form.get('draw'))
    start = int(request.form.get('start'))
    length = int(request.form.get('length'))
    search_value = request.form.get('search[value]')

    total_records, filtered_records, result = retrieve_all_employee_data(search_value,length,start)

    return jsonify({
        'draw': draw,
        'recordsTotal': total_records,
        'recordsFiltered': filtered_records,
        'data': result
    })

@app.route('/manage_cache_employee')
@personal_access_required
def manage_cache_employee():
    data = fetch_all_cache_employee_data()
    return render_template('cache_employees.html', data=data)

@app.route('/delete_buyer', methods=['POST'])
@personal_access_required
def delete_buyer():
    data = request.get_json()
    buyer_id = data['buyer_id']
    ip_address = request.remote_addr
    username = session['email']
    if delete_buyers(buyer_id):
        insert_log_to_db(str(buyer_id),'Personal Account','Buyer Deleted',ip_address,username,'','')
        return jsonify({'message': 'Data Deleted successfully'}), 200
    else:
        return jsonify({'message': 'Error updating data'}), 500


@app.route('/add_buyer', methods=['POST'])
@personal_access_required
def add_buyer():
    new_buyer_name = request.form['buyer_name']
    new_buyer_logo = request.files['buyer_logo']
    if new_buyer_name == 'null':
        new_buyer_name = None
    # Check if the buyer_logo field exists and the file is valid
    if new_buyer_logo:
        upload_folder = r"static/images/buyer_logos"
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder)

        # Update the filename pattern to include the correct extension
        filename = f"{new_buyer_name}.{new_buyer_logo.filename.split('.')[-1]}"
        filename = filename.replace(' ', '_')
        image_path = os.path.join(upload_folder, filename)

        try:
            new_buyer_logo.save(image_path)
            print('Image saved:', image_path)
        except Exception as e:
            print('Error saving image:', e)

    if add_new_buyer(new_buyer_name):
        ip_address = request.remote_addr
        username = session['email']
        insert_log_to_db('','Personal Account','Buyer Added',ip_address,username,'',new_buyer_name)
        return jsonify({'message': 'Data Added successfully'}), 200
    else:
        return jsonify({'message': 'Error updating data'}), 500




@app.route('/update_buyer', methods=['POST'])
@personal_access_required
def update_buyer():
    data = request.get_json()
    buyer_id = data['buyer_id']
    edited_buyer_name = data['buyer_name']
    username = data['username']
    password = data['password']

    if update_buyer_name(buyer_id,edited_buyer_name,username,password):
        ip_address = request.remote_addr
        username = session['email']
        insert_log_to_db(str(buyer_id),'Personal Account','Buyer Updated',ip_address,username,'',edited_buyer_name)
        return jsonify({'message': 'Data updated successfully'}),200
    else:
        return jsonify({'message': 'Error updating data'}), 500


@app.route('/manage_buyer')
@personal_access_required
def manage_buyer():
    data = get_all_buyers_data()
    company_names = get_all_company_names()
    unique_buyers = set()
    filtered_data = []

    for buyer in data:
        if buyer['buyer_name'] not in unique_buyers:
            unique_buyers.add(buyer['buyer_name'])
            filtered_data.append(buyer)


    return render_template('buyer_crud.html', data=data,company_names=company_names,buyer_names =filtered_data)

@app.route('/manage_offices')
@personal_access_required
def manage_offices():
    data = get_all_offices_data()
    company_names = get_all_company_names()
    print('offices',data)
    print('offices',company_names)
    return render_template('office_crud.html', data=data, company_names=company_names)

@app.route('/logout')
@handle_session_error
def logout():
    ip_address = request.remote_addr
    username = session['email']
    access_id = session['access_id']
    insert_log_to_db(str(access_id),'Logged Out','Logged Out',ip_address,username,'Logged Out','')

    return redirect(url_for('login'))

@app.route('/update_employee', methods=['POST'])
@personal_access_required
def update_employee():
    edited_data = request.json
    ip_address = request.remote_addr
    username = session['email']
    if update_employee_data(edited_data):
        insert_log_to_db('','Personal Account','Employee Updated',ip_address,username,edited_data,'')
        return jsonify({'message': 'Data updated successfully'}),200
    else:
        return jsonify({'message': 'Error updating data'}), 500
@app.route('/add_buyer_company', methods=['POST'])
@personal_access_required
def update_buyer_company():
    edited_data = request.json
    ip_address = request.remote_addr
    username = session['email']
    if add_buyer_company(edited_data['company_name'],edited_data['buyer_name']):
        insert_log_to_db('','Personal Account','New Company Linked',ip_address,username,f'{edited_data["company_name"]} - {edited_data["buyer_name"]}','')
        return jsonify({'message': 'Data updated successfully'})
    else:
        return jsonify({'message': 'Error updating data'}), 500

@app.route('/generate_cards', methods=['POST'])
@personal_access_required
def generate_cards():
    file = request.files['file']

    if file and allowed_file(file.filename):
        try:
            if file.filename.endswith('.xlsx'):
                df = pd.read_excel(file)
            elif file.filename.endswith('.csv'):
                df = pd.read_csv(file)
            else:
                return jsonify({'error': 'Unsupported file format. Please upload an Excel (xlsx) or CSV file.'})
            # Queue to store employee IDs for image generation
            employee_ids_for_qr = []
    
            for index, row in df.iterrows():
                employee_data = {
                    'company_id': row['company_id'],
                    'employee_id': row['employee_id']
                }
                employee_ids_for_qr.append(employee_data['employee_id'])
                generate_and_save_qr_image(employee_data['employee_id'],employee_data['company_id'])

            
            source_dir = '/home/lcxacjx9hurd/public_html/complaintSystem/static/images/card_images'
            zip_filename = '/home/lcxacjx9hurd/public_html/complaintSystem/static/images/card_images.zip'
    
            zipf = zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED)
    
            # Walk through the 'card_images' directory and add all files to the zip
            for root, _, files in os.walk(source_dir):
                for file in files:
                    zipf.write(os.path.join(root, file), os.path.relpath(os.path.join(root, file), source_dir))
    
            zipf.close()

    
            return jsonify({'message': 'File uploaded and employees added successfully'}), 200
        except Exception as e:
            print('Error',e)
            return jsonify({'error': 'Invalid file'}), 500
    return jsonify({'error': 'Invalid file'}), 500
@app.route('/upload_employee_file', methods=['POST'])
@personal_access_required
def upload_employee_file():
    file = request.files['file']
    filename = file.filename
    ip_address = request.remote_addr
    username = session['email']
    if file and allowed_file(file.filename):
        if filename.endswith('.xlsx'):
            df = pd.read_excel(file, dtype={'mobile_number': str, 'cnic_no': str})
        elif filename.endswith('.csv'):
            df = pd.read_csv(file, dtype={'mobile_number': str, 'cnic_no': str})
        else:
            return jsonify({'error': 'Unsupported file format. Please upload an Excel (xlsx) or CSV file.'})
        # Queue to store employee IDs for image generation
        employee_ids_for_qr = []
        insert_log_to_db('','Personal Account','Employee Data Uploaded',ip_address,username,filename,'')
        for index, row in df.iterrows():
            if index == 0:
                last_employee_id = get_last_fos_id()
            
            employee_data = {
                'employee_name': row['employee_name'],
                'worker_type': row['worker_type'],
                'department': row['department'],
                'designation': row['designation'],
                'mobile_number': row['mobile_number'],
                'gender': row['gender'],
                'office_id': row['office_id'],
                'cnic_no': row['cnic_no'],
                'employee_id': generate_employee_ids(row,last_employee_id),
                'company_id': row['company_id']
            }
            try:
                employee_data['location'] = row['temp_data']
            except:
                employee_data['location'] = ''
            if add_employee_to_database(employee_data):
                last_employee_id = employee_data['employee_id']
                employee_ids_for_qr.append(employee_data['employee_id'])
            else:
                return jsonify({'error': 'Failed to add an employee to the database'}), 500

        return jsonify({'message': 'File uploaded and employees added successfully'}), 200
    return jsonify({'error': 'Invalid file'}), 500


@app.route('/delete_employee', methods=['POST'])
@personal_access_required
def delete_employee_id():
    try:
        data = request.json
        employee_id = data.get('employee_id')
        ip_address = request.remote_addr
        username = session['email']
        delete_employee(employee_id)
        insert_log_to_db('','Personal Account','Employee Deleted',ip_address,username,employee_id,'')
        return jsonify({'message': 'Employee deleted successfully'}), 200
    except Exception as e:
        print("Error deleting employee:", e)
        return jsonify({'error': 'Invalid Employee'}), 500
@app.route('/check_status')
def check_status():
    return render_template('complaint_timeline.html')
@app.route('/search_complaint', methods=['POST'])
def search_complaint():
    data = request.get_json()
    search_term = data.get('query')
    data = []
    specificComplaint = get_specific_complaint_data(search_term)
    if specificComplaint:
        return jsonify(specificComplaint)
        data.append(specificComplaint)
        #specificComplaint = load_complaints_images(data)
        specificComplaint = specificComplaint[0]
        print(specificComplaint)
    else:
        specificComplaint = None
    return jsonify(specificComplaint)

@app.route('/employee/<int:employee_id>')
def employee_profile(employee_id):
    # Fetch employee data from database based on employee_id
    employee_data = get_employee_data_for_fos_card(employee_id)
    if employee_data:
        return render_template('employee_profile.html', employee=employee_data)
    else:
        return "Employee not found", 404

@app.route('/update_company', methods=['POST'])
@personal_access_required
def update_company():
    edited_data = request.json
    ip_address = request.remote_addr
    username = session['email']
    if update_company_data(edited_data):
        insert_log_to_db('','Personal Account','Company Updated',ip_address,username,edited_data,'')
        return jsonify({'message': 'Company data updated successfully'})
    else:
        return jsonify({'message': 'Error updating company data'}), 500

@app.route('/delete_company', methods=['POST'])
@personal_access_required
def delete_company():
    try:
        data = request.json
        company_id = data.get('company_id')
        delete_company_data(company_id)
        ip_address = request.remote_addr
        username = session['email']  # Implement the function to delete a company
        insert_log_to_db('','Personal Account','Company Deleted',ip_address,username,company_id,'')
        return jsonify({'message': 'Company deleted successfully'}), 200
    except Exception as e:
        print("Error deleting company:", e)
        return jsonify({'error': 'Invalid Company'}), 500


@app.route('/update_office', methods=['POST'])
@personal_access_required
def update_office():
    data = request.json
    ip_address = request.remote_addr
    username = session['email'] 
    if update_office_data(data):
        insert_log_to_db('','Personal Account','Office Updated',ip_address,username,data,'')
        return jsonify({'message': 'Office data updated successfully'}),200
    else:
        return jsonify({'message': 'Error updating company data'}), 500
@app.route('/add_office', methods=['POST'])
@personal_access_required
def add_office():
    data = request.get_json()
    office_name = data['office_name']
    office_location = data['office_location']
    company_name = data['company_name']
    mobile_number = data['mobile_number']
    ip_address = request.remote_addr
    username = session['email']
    if add_new_office(office_name,office_location, company_name,mobile_number):
        insert_log_to_db('','Personal Account','Office Added',ip_address,username,data,'')
        return jsonify({'message': 'Office added successfully'}), 200
    else:
        return jsonify({'error': 'Failed to add office'}), 500

@app.route('/delete_office', methods=['POST'])
@personal_access_required
def delete_office():

    data = request.get_json()
    office_id = data.get('office_id')
    ip_address = request.remote_addr
    username = session['email']
    if delete_office_from_database(office_id):
        insert_log_to_db('','Personal Account','Office Deleted',ip_address,username,office_id,'')
        response = {'message': 'Office deleted successfully'}
        return jsonify(response), 200
    else:
        response = {'error': 'Failed to delete office'}
        return jsonify(response), 500

@app.route('/add_company', methods=['POST'])
@personal_access_required
def add_company():
    try:
        data = request.form
        company_name = data.get('name')
        print(data)
        ip_address = request.remote_addr
        username = session['email']
        
        # Check if a logo file was provided
        if 'logo' in request.files and request.files['logo'].filename != '':
            new_company_logo = request.files['logo']
            upload_folder = r"static/images/company_logos"
            if not os.path.exists(upload_folder):
                os.makedirs(upload_folder)
            # Update the filename pattern to include the correct extension
            filename = f"{company_name}.{new_company_logo.filename.split('.')[-1]}"
            filename = filename.replace(' ', '_')
            image_path = os.path.join(upload_folder, filename)
            try:
                new_company_logo.save(image_path)
                print('Image saved:', image_path)
            except Exception as e:
                print('Error saving image:', e)
        else:
            print('No logo provided')
        
        if add_new_company(company_name):
            insert_log_to_db('','Personal Account','Company Added',ip_address,username,company_name,'')
            return jsonify({'message': 'Company added successfully'}), 200
        else:
            return jsonify({'message': 'Error updating company data'}), 500
    except Exception as e:
        print("Error adding company:", e)
        return jsonify({'error': 'Failed to add company'}), 500


def load_complaints_images(fetched_complaints):
    # Get a list of CAPA image filenames from the static/images/capa_images folder
    capa_images_folder = os.path.join(app.root_path, 'static', 'images', 'capa_images')
    capa_image_filenames = os.listdir(capa_images_folder)

    # Iterate through fetched_complaints and associate CAPA image URLs
    for complaint in fetched_complaints:
        ticket_number = complaint["ticket_number"]

        temp_path = f'{ticket_number}_capa.jpg'
        temp_path1 = f'{ticket_number}_capa.png'
        temp_path2 = f'{ticket_number}_capa.jpeg'

        if temp_path.lower() in [x.lower() for x in capa_image_filenames]:
            complaint['capaImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path}')
        elif temp_path1 in capa_image_filenames:
            complaint['capaImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path1}')
        elif temp_path2 in capa_image_filenames:
            complaint['capaImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path2}')
        else:
            complaint['capaImageURL'] = ''

        temp_path = f'{ticket_number}_capa1.jpg'
        temp_path1 = f'{ticket_number}_capa1.png'
        temp_path2 = f'{ticket_number}_capa1.jpeg'

        if temp_path.lower() in [x.lower() for x in capa_image_filenames]:
            complaint['capa1ImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path}')
        elif temp_path1 in capa_image_filenames:
            complaint['capa1ImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path1}')
        elif temp_path2 in capa_image_filenames:
            complaint['capa1ImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path2}')
        else:
            complaint['capa1ImageURL'] = ''

        temp_path = f'{ticket_number}_capa2.jpg'
        temp_path1 = f'{ticket_number}_capa2.png'
        temp_path2 = f'{ticket_number}_capa2.jpeg'

        if temp_path.lower() in [x.lower() for x in capa_image_filenames]:
            complaint['capa2ImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path}')
        elif temp_path1 in capa_image_filenames:
            complaint['capa2ImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path1}')
        elif temp_path2 in capa_image_filenames:
            complaint['capa2ImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path2}')
        else:
            complaint['capa2ImageURL'] = ''
        temp_path = f'{ticket_number}_proof.jpg'
        temp_path1 = f'{ticket_number}_proof.png'
        temp_path2 = f'{ticket_number}_proof.jpeg'
        
        if temp_path.lower() in [x.lower() for x in capa_image_filenames]:
            complaint['proofImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path}')
        elif temp_path1 in capa_image_filenames:
            complaint['proofImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path1}')
        elif temp_path2 in capa_image_filenames:
            complaint['proofImageURL'] = url_for('static', filename=f'images/capa_images/{temp_path2}')
        else:
            complaint['proofImageURL'] = ''
    return fetched_complaints



def get_company_logo(buyer_name):
    # Get a list of CAPA image filenames from the static/images/capa_images folder
    buyer_name = buyer_name.replace(' ', '_')
    buyer_images_folder = os.path.join(app.root_path, 'static', 'images', 'buyer_logos')
    buyer_images_filenames = os.listdir(buyer_images_folder)
    company_images_folder = os.path.join(app.root_path, 'static', 'images', 'company_logos')
    company_images_filenames = os.listdir(company_images_folder)
    for image in buyer_images_filenames:
        if buyer_name in image:
            image = image.replace(' ', '_')
            return url_for('static', filename='images/buyer_logos/{}'.format(image))
        else:
            for company_image in company_images_filenames:
                if buyer_name in company_image:
                    company_image = company_image.replace(' ', '_')
                    return url_for('static', filename='images/company_logos/{}'.format(company_image))

    return None
def get_company_logo_for_report(buyer_name):
    # Get a list of CAPA image filenames from the static/images/capa_images folder
    buyer_name = buyer_name.replace(' ', '_')

    buyer_images_folder = os.path.join(app.root_path, 'static', 'images', 'buyer_logos')
    buyer_images_filenames = os.listdir(buyer_images_folder)
    company_images_folder = os.path.join(app.root_path, 'static', 'images', 'company_logos')
    company_images_filenames = os.listdir(company_images_folder)
    
    for image in buyer_images_filenames:
        if buyer_name in image:
            image = image.replace(' ', '_')
            return 'static/images/buyer_logos/{}'.format(image)
        else:
            for company_image in company_images_filenames:
                if buyer_name in company_image:
                    company_image = company_image.replace(' ', '_')
                    return 'static/images/company_logos/{}'.format(company_image)

    return None



def load_companies_images(fetched_complaints):
    # Get a list of CAPA image filenames from the static/images/capa_images folder
    capa_images_folder = os.path.join(app.root_path, 'static', 'images', 'company_logos')
    capa_image_filenames = os.listdir(capa_images_folder)

    # Create a dictionary to store CAPA image URLs with ticket numbers
    images_dict = {}
    for key in fetched_complaints:
        formatted_key = key.replace(' ', '_')
        images_dict[key] = ''
        for image in capa_image_filenames:
            if formatted_key in image:
                images_dict[key] = url_for('static', filename='images/company_logos/{}'.format(image))

    return images_dict


import os

def save_uploaded_image(ticket_number, file, capa_identifier):
    if file:
        MAX_FILE_SIZE_MB = 5
        upload_folder = "static/images/capa_images"
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder)

        # Use secure_filename to ensure a safe filename
        filename = "{}_{}".format(ticket_number, capa_identifier)
        # filename = secure_filename(filename)

        # Determine file extension
        _, file_extension = os.path.splitext(file.filename)

        # Check file size
        max_file_size_bytes = MAX_FILE_SIZE_MB * 1024 * 1024  # Convert MB to bytes
        if file.content_length > max_file_size_bytes:
            print('File size exceeds the limit ({} MB)'.format(MAX_FILE_SIZE_MB))
            return

        # Check if the file is closed
        if not file.closed:
            # Reset file position to the beginning
            file.seek(0)

            # Save file based on its type
            if file_extension.lower() == ".pdf":
                save_as_pdf(file, upload_folder, filename)
            elif file_extension.lower() in (".mp4", ".avi", ".mkv", ".mov"):
                save_as_video(file, upload_folder, filename)
            elif file_extension.lower() in (".jpg", ".jpeg", ".png", ".gif"):
                save_as_image(file, upload_folder, filename)
            else:
                print('Unsupported file type:', file.filename)

            # Close the file after saving
            file.close()
        else:
            print('File is already closed')

        
        
def save_as_pdf(file, upload_folder, filename):
    pdf_path = os.path.join(upload_folder, filename + ".pdf")
    try:
        file.save(pdf_path)
        print('PDF saved:', pdf_path)
    except Exception as e:
        print('Error saving PDF:', e)

def save_as_video(file, upload_folder, filename):
    video_path = os.path.join(upload_folder, filename + ".mp4")
    try:
        file.save(video_path)
        print('Video saved:', video_path)
    except Exception as e:
        print('Error saving video:', e)

def save_as_image(file, upload_folder, filename):
    image_path = os.path.join(upload_folder, filename + ".png")
    try:
        file.save(image_path)
        print('Image saved:', image_path)
    except Exception as e:
        print('Error saving image:', e)

#File Format
#employee_id	employee_name	worker_type	department	designation	mobile_number	gender	office_id	cnic_no	
@app.route('/add_manual', methods=['POST'])
def add_manual():
    try:
        file = request.files['file']
        filename = file.filename
        print(f"Received file for Add Manually: {filename}")

        if filename.endswith('.xlsx'):
            employee_data = pd.read_excel(file, dtype={'mobile_number': str, 'cnic_no': str})
        elif filename.endswith('.csv'):
            employee_data = pd.read_csv(file, dtype={'mobile_number': str, 'cnic_no': str})
        else:
            return jsonify({'error': 'Unsupported file format. Please upload an Excel (xlsx) or CSV file.'})
        
        # Insert each employee into the database
        for index, row in employee_data.iterrows():
            employee_data = {
                'employee_name': row['employee_name'],
                'worker_type': row['worker_type'],
                'department': row['department'],
                'designation': row['designation'],
                'mobile_number': row['mobile_number'],
                'gender': row['gender'],
                'office_id': row['office_id'],
                'cnic_no': row['cnic_no'],
                'employee_id': row['employee_id'],
                'location': row['temp_data'],
                'company_id': row['company_id']
            }
            if add_employee_to_database(employee_data):
                #employee_ids_for_qr.append(employee_data['employee_id'])
                pass
            else:
                return jsonify({'error': 'Failed to add an employee to the database'}), 500

        return jsonify({'message': 'Data inserted into the database successfully!'}), 200
    except Exception as e:
        print(e)
        return jsonify({'error': str(e)}), 500

#File Format
#employee_name	worker_type	department	designation	mobile_number	gender	office_id	cnic_no	
@app.route('/download_leavers', methods=['POST'])
def download_leavers():
    try:
        file = request.files['file']
        filename = file.filename
        print(f"Received file for Download Joiners/Leavers: {filename}")

        # Convert file to DataFrame
        generate_leavers_joiners(file)

        # Return the files as attachments
        file_paths = ['dec_leavers_df.xlsx', 'dec_joiners_df.xlsx', 'updated_employee_data.xlsx','transfer_cases.xlsx']

        # Create a BytesIO object to hold the zip file in memory
        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
            for file_path in file_paths:
                zip_file.write(file_path, os.path.basename(file_path))

        # Move the buffer's position to the beginning to ensure all data is written
        zip_buffer.seek(0)

        return send_file(
            zip_buffer,
            download_name='files.zip',
            as_attachment=True
        )

    except Exception as e:
        print(str(e))
        return jsonify({'error': str(e)})
        
#File Format
#employee_id	employee_name	worker_type	department	designation	mobile_number	gender	office_id	cnic_no	
@app.route('/update_employees', methods=['POST'])
def update_employees():
    try:
        file = request.files['file']
        filename = file.filename
        print(f"Received file for Update Employees: {filename}")
        ip_address = request.remote_addr
        username = session['email']
        if filename.endswith('.xlsx'):
            df = pd.read_excel(file)
        elif filename.endswith('.csv'):
            df = pd.read_csv(file)
        else:
            return jsonify({'error': 'Unsupported file format. Please upload an Excel (xlsx) or CSV file.'})
        insert_log_to_db('','Personal Account','Employee Data Updated',ip_address,username,filename,'')
        for index, row in df.iterrows():
            if index == 0:
                last_employee_id = get_last_fos_id()
            employee_data = {
                'employee_name': row['employee_name'],
                'worker_type': row['worker_type'],
                'department': row['department'],
                'designation': row['designation'],
                'mobile_number': row['mobile_number'],
                'gender': row['gender'],
                'office_id': row['office_id'],
                'cnic_no': row['cnic_no'],
                'company_id': row['company_id']
            }

            if is_employee_present(employee_data['cnic_no']):
                update_employee_in_database(employee_data)
            else:
                employee_data['employee_id'] = generate_employee_ids(row,last_employee_id)
                add_employee_to_database(employee_data)

        return jsonify({'message': 'Data updated/inserted successfully!'})

    except Exception as e:
        print(str(e))
        return jsonify({'error': str(e)})


#File Format
#employee_id	employee_name	worker_type	department	designation	mobile_number	gender	office_id	cnic_no	
@app.route('/add_leavers', methods=['POST'])
def add_leavers():
    try:
        file = request.files['file']
        filename = file.filename
        print(f"Received file for Add Leavers: {filename}")
        ip_address = request.remote_addr
        username = session['email']
        insert_log_to_db('','Personal Account','Added Leavers',ip_address,username,filename,'')
        # Read the file into a DataFrame
        if filename.endswith('.xlsx'):
            employee_data = pd.read_excel(file)
        elif filename.endswith('.csv'):
            employee_data = pd.read_csv(file)
        else:
            return jsonify({'error': 'Unsupported file format. Please upload an Excel (xlsx) or CSV file.'})
        
        # Extract employee_ids from the file
        employee_ids = employee_data['employee_id'].tolist()
        print(employee_ids)
        # Update the employee_left column in the database
        update_employee_left_in_database(employee_ids)

        # Process the file as needed
        # You can access file data using file.read(), file.filename, etc.

        return jsonify({'message': 'File uploaded successfully!'})
    except Exception as e:
        print(str(e))
        return jsonify({'error': str(e)})
@app.route('/approve_complaint_request', methods=['POST'])
@cs_access_required
def approve_complaint_request():
    try:
        data = request.json
        
        ticket_number = data.get('ticket')
        mobile_number = data.get('mobile')
        office_id = data.get('office')
        employee_id = data.get('employee')
        is_anonymous = data.get('anonymous')
        category = data.get('category')
        complaint_text = data.get('complaint_text')
        print('ticket',ticket_number, 'mobile',mobile_number, 'office',office_id, 'employee',employee_id, 'anonymous',is_anonymous, 'category',category)
        print('ticket_number',ticket_number)
        if approve_complaint(ticket_number):
            ip_address = request.remote_addr
            username = session['email']
            insert_log_to_db('','CS Portal','Complaint Approved',ip_address,username,'',ticket_number)
            io_mobile_number,office_id = get_io_phone_no(employee_id,category,complaint_text)
            if is_anonymous == True or employee_id == 0:
                message = f"New Complaint #{ticket_number} is registered from an anonymous employee."
            else:
                message = f"New Complaint #{ticket_number} is registered from employee #{employee_id}."

            add_notifications(message,office_id)
            send_sms_to_io(io_mobile_number, ticket_number)
            send_notification(employee_id,ticket_number)
            send_sms(mobile_number, ticket_number)

            response_data = {'message': 'Complaint approved successfully'}
            return jsonify(response_data), 200
        else:
            response_data = {'message': 'Complaint could not be approved'}
            return jsonify(response_data), 400  # You can choose an appropriate status code here
    except Exception as e:
        # Handle any errors that might occur during processing
        error_message = {'error': str(e)}
        return jsonify(error_message), 500

@app.route('/reject_complaint_request', methods=['POST'])
@cs_access_required
def reject_complaint_request():
    try:
        data = request.json
        
        ticket_number = data.get('ticket')
        ip_address = request.remote_addr
        username = session['email']
        if reject_complaint(ticket_number):
            insert_log_to_db('','CS Portal','Complaint Rejected',ip_address,username,'',ticket_number)
            response_data = {'message': 'Complaint rejected successfully'}
            return jsonify(response_data), 200
        else:
            response_data = {'message': 'Complaint could not be rejected'}
            return jsonify(response_data), 400  # You can choose an appropriate status code here
    except Exception as e:
        # Handle any errors that might occur during processing
        error_message = {'error': str(e)}
        return jsonify(error_message), 500
@app.route('/send_complaint_otp', methods=['POST'])
def send_complaint_otp():
    try:
        data = request.get_json()
        mobile_number = data.get('mobile_number')
        otp = generate_otp()
        print('OTP Verification', otp)
        success = send_login_otp(otp, mobile_number)
        
        if success:
            return jsonify({'message': 'OTP sent successfully','otp': otp}),200
        else:
            return jsonify({'message': 'Failed to send OTP'}), 500

    except Exception as e:
        print('An error occurred:', str(e))
        return jsonify({'message': 'Internal server error'}),400
@app.route('/survey', methods=['GET', 'POST'])
@personal_access_required
def survey():
    surveys = get_all_surveys()
    offices = fetch_all_offices()
    departments = fetch_all_departments()  # New function to fetch departments
    return render_template('survey_crud.html', data=surveys, offices=offices, departments=departments)
@app.route('/launch_survey_crud', methods=['GET', 'POST'])
@admin_access_required
def launch_survey_crud():
    access_id = session['access_id']
    surveys = get_all_surveys(access_id)
    offices = fetch_all_offices(access_id)
    departments = fetch_all_departments(access_id)
    return render_template('launchSurvey.html', data=surveys, offices=offices, departments=departments)

@app.route('/add_survey', methods=['POST'])
def add_survey():
    data = request.get_json()
    access_id = session['access_id']
    new_survey = {
        "title": data.get('title'),
        "description": data.get('description'),
        "question_count": data.get('question_count'),
        "estimated_time": data.get('estimated_time'),
        "created_at": datetime.now().strftime("%Y-%m-%d"),
        "expiry_date": data.get('date_expiry'),
        "filters": data.get('filters')
    }
    print('New Survey:', new_survey)
    response = add_new_survey(new_survey, access_id)
    return jsonify(response), 201

@app.route('/edit_survey', methods=['POST'])
def edit_survey():
    data = request.json
    if edit_survey_in_db(data['surveyId'], data):
        return jsonify({"message": "Survey updated successfully"}), 200
    else:
        return jsonify({"error": "Failed to update survey"}), 500

@app.route('/delete_survey', methods=['POST'])
def delete_survey():
    survey_id = request.json['surveyId']
    success, message = delete_survey_from_db(survey_id)
    if success:
        return jsonify({"message": message}), 200
    else:
        return jsonify({"error": message}), 500
@app.route('/logout_user', methods=['POST'])
def logout_user():
    ip_address = request.remote_addr
    username = session['email']
    session_id = request.form.get('session_id')
    logout_user(session_id)
    insert_log_to_db('','Logout','User Logged Out',ip_address,username,'',session_id)
    return redirect(url_for('show_logged_users'))
@app.route('/launch_survey_questions', methods=['GET'])
@admin_access_required
def launch_survey_questions():
    access_id = session['access_id']
    questions = get_all_survey_questions(access_id)  # Implement this function to fetch all questions
    surveys = get_all_surveys_for_crud(access_id)  # Implement this function to fetch all surveys
    return render_template('launch_survey_questions.html', questions=questions, surveys=surveys)


@app.route('/survey_questions', methods=['GET'])
@personal_access_required
def survey_questions():
    questions = get_all_survey_questions()  # Implement this function to fetch all questions
    surveys = get_all_surveys_for_crud()  # Implement this function to fetch all surveys
    return render_template('survey_questions_crud.html', questions=questions, surveys=surveys)

@app.route('/add_question', methods=['POST'])
def add_question():
    data = request.form
    instruction_media = request.files.get('instruction_media')
    
    success, message = add_question_to_db(data,instruction_media)
    if success:
        return jsonify({"message": message}), 200
    else:
        return jsonify({"error": message}), 500

@app.route('/edit_question', methods=['POST'])
def edit_question():
    data = request.json
    success, message = update_question_in_db(data)
    if success:
        return jsonify({"message": message}), 200
    else:
        return jsonify({"error": message}), 500

@app.route('/delete_question', methods=['POST'])
def delete_question():
    data = request.json
    success, message = delete_question_from_db(data['question_id'])
    if success:
        return jsonify({"message": message}), 200
    else:
        return jsonify({"error": message}), 500
def send_email(recipient_email, username, password):
    smtp_server = "162.240.164.167"
    port = 465  # For unencrypted/TLS connections
    sender_email = "ilessdb@ilessdb.org"
    sender_password = "M.m03007493358"

    message = MIMEMultipart("alternative")
    message["Subject"] = "Your ILESSDB Dashboard Credentials"
    message["From"] = sender_email
    message["To"] = recipient_email

    # Create the HTML version of your message
    html = f"""
    <html>
      <body>
        <h2>Welcome to ILESSDB Dashboard</h2>
        <p>Here are your login credentials:</p>
        <p><strong>Username:</strong> {username}</p>
        <p><strong>Password:</strong> {password}</p>
        <p>Please change your password upon first login.</p>
        <p>If you have any questions, please contact our support team.</p>
      </body>
    </html>
    """

    # Turn these into plain/html MIMEText objects
    part = MIMEText(html, "html")

    # Add HTML part to MIMEMultipart message
    message.attach(part)

    try:
        # Create a secure SSL/TLS context
        context = smtplib.ssl.create_default_context()
        context.check_hostname = False
        context.verify_mode = smtplib.ssl.CERT_NONE

        # Try to log in to server and send email
        with smtplib.SMTP(smtp_server, port) as server:
            server.set_debuglevel(2)  # Enable verbose debug output
            server.ehlo()  # Can be omitted
            server.starttls(context=context)
            server.ehlo()  # Can be omitted
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipient_email, message.as_string())
        return {"success": True, "message": "Email sent successfully"}
    except Exception as e:
        print(f"Email could not be sent. Error: {e}")
        return {"success": False, "message": f"Email could not be sent. Error: {str(e)}"}

last_email_time = 0

@app.route('/send-email', methods=['POST'])
def handle_send_email():
    global last_email_time
    data = request.json
    recipient_email = data.get('email')
    username = data.get('username')
    password = data.get('password')

    # Input validation
    if not recipient_email or '@' not in recipient_email:
        return jsonify({"success": False, "message": "Invalid email address"}), 400

    if not username or not password:
        return jsonify({"success": False, "message": "Username and password are required"}), 400

    # Rate limiting
    current_time = time.time()
    if current_time - last_email_time < 60:  # 60 seconds
        return jsonify({"success": False, "message": "Please wait before sending another email"}), 429
    last_email_time = current_time

    result = send_email(recipient_email, username, password)
    return jsonify(result)
@app.route('/chat')
def chat():
    return render_template('chatbot.html')

@app.route('/ask', methods=['POST'])
def ask():
    user_message = request.json['message']
    response = ask_question(user_message)
    return jsonify({'message': response})
@app.route('/send_factory_form', methods=['POST'])
def send_factory_form():
    if request.method == 'POST':
        form_data = {}
        for field, value in request.form.items():
            if value:  # Only include non-empty fields
                form_data[field] = value
        
        # Here you can add additional processing, validation, etc.
        
        try:
            send_factory_form_email(form_data)
            print('Form submitted successfully!', 'success')
        except Exception as e:
            print(f'An error occurred: {str(e)}', 'error')
        
        return {"success": True, "message": "Email sent successfully"}
@app.route('/send_contact_form', methods=['POST'])
def send_contact_form():
    if request.method == 'POST':
        form_data = {}
        for field, value in request.form.items():
            if value:  # Only include non-empty fields
                form_data[field] = value
        
        # Here you can add additional processing, validation, etc.
        
        try:
            send_contact_form_email(form_data)
            print('Form submitted successfully!', 'success')
        except Exception as e:
            print(f'An error occurred: {str(e)}', 'error')
        
        return {"success": True, "message": "Email sent successfully"}
@app.route('/survey_search')
def survey_search():
    return render_template('survey_search.html')
@app.route('/survey_list')
def survey_list():
    employee_id = request.args.get('employee_id')
    if not employee_id:
        return "No employee ID provided", 400
    
    print('Employee ID:', employee_id)
    # Perform any additional server-side verification if needed
    return render_template('survey_list.html', employee_id=employee_id)
@app.route('/survey_questions_form')
def survey_questions_form():
    employee_id = request.args.get('employeeId')
    survey_id = request.args.get('surveyId')
    if not employee_id and not survey_id:
        return "No employee ID provided", 400
    
    # Perform any additional server-side verification if needed
    return render_template('survey_questions.html', employee_id=employee_id,survey_id = survey_id)
def getFeedbackComplaintsSummary(complaints):
    result = []
    for complaint in complaints:
        ticket_number = complaint['ticket_number']
        complaint_text = complaint['additional_comments']
        category = complaint['complaint_categories']
        entry_date = complaint['date_entry']
        if category == 'Feedback':
            try:
                summary = get_or_generate_summary(ticket_number, complaint_text)
                result.append({
                    'ticket_number': ticket_number,
                    'summary': summary,
                    'entry_date': entry_date
                })
            except Exception as e:
                print(f"Error processing complaint {ticket_number}: {str(e)}")
    return result
def getDormitoryComplaintsSummary(complaints):
    result = []
    for complaint in complaints:
        ticket_number = complaint['ticket_number']
        complaint_text = complaint['additional_comments']
        category = complaint['complaint_categories']
        entry_date = complaint['date_entry']
        
        # Check if complaint text contains 'dormitory complaint' case-insensitive
        if 'dormitory complaint' in complaint_text.lower():
            try:
                summary = generate_summary(complaint_text)
                result.append({
                    'ticket_number': ticket_number,
                    'summary': summary,
                    'entry_date': entry_date
                })
            except Exception as e:
                print(f"Error generating summary for complaint {ticket_number}: {str(e)}")
    return result
def fetch_and_update_employees():
    print("Starting the employee data update procedure...")
    
    base_url = "https://jj-api.resourceinn.com"
    client_id = "2"
    client_secret = "rrCIFk1OhWVQzeb0jAKWpLuCzFtJND9VpKrD3X6M"
    username = "HRMANAGER"
    password = "Junaid007"

    db_config = {
        'host': 'localhost',
        'user': 'fos-database',
        'password': 'M.m03007493358',
        'database': 'fos-database'
    }

    branch_office_id = {
        'CC': 181,
        'PHASE-6': 182,
        'CLOUD KITCHEN': 183,
        'EMPORIUM': 184,
        'JJ BAHRIA TOWN': 185,
        'JT': 186,
        'WEHSHI LAB RAYA P6': 187,
    }

    insert_employees_from_api(base_url, client_id, client_secret, username, password, db_config, branch_office_id)
class SingletonScheduler:
    _instance = None
    _lock = threading.Lock()
    
    @classmethod
    def get_instance(cls):
        if cls._instance is None:
            with cls._lock:
                if cls._instance is None:
                    # Configure logging for the scheduler
                    logging.basicConfig(
                        level=logging.INFO,
                        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
                    )
                    scheduler_logger = logging.getLogger('apscheduler')
                    scheduler_logger.setLevel(logging.INFO)
                    
                    # Create scheduler with more explicit configuration
                    cls._instance = BackgroundScheduler(
                        timezone=pytz.timezone('Asia/Karachi'),
                        job_defaults={
                            'coalesce': True,  # Combine multiple missed runs into one
                            'max_instances': 1,  # Prevent multiple instances of same job running
                            'misfire_grace_time': 15 * 60  # 15 minutes grace time for misfires
                        }
                    )
                    
                    # Add event listeners for job execution tracking
                    cls._instance.add_listener(
                        cls._job_execution_listener,
                        EVENT_JOB_EXECUTED | EVENT_JOB_ERROR
                    )
        
        return cls._instance
    
    @staticmethod
    def _job_execution_listener(event):
        if event.exception:
            logging.error(f'Job {event.job_id} failed: {str(event.exception)}')
        else:
            logging.info(f'Job {event.job_id} completed successfully')



def print_scheduler_jobs():
    scheduler = SingletonScheduler.get_instance()
    logging.info("Current scheduled jobs:")
    for job in scheduler.get_jobs():
        next_run = job.next_run_time.strftime("%Y-%m-%d %H:%M:%S %Z")
        logging.info(f"Job: {job.id}, Function: {job.func.__name__}, Next run: {next_run}")

def init_scheduler():
    scheduler = SingletonScheduler.get_instance()
    
    if not scheduler.running:
        try:
            # Add jobs with IDs for better tracking
            scheduler.add_job(
                fetch_and_update_employees,
                'cron',
                hour=22,
                minute=7,
                id='fetch_employees_job',
                replace_existing=True
            )
            
            scheduler.add_job(
                add_notifications_for_late_capas,
                'cron',
                hour=8,
                minute=0,
                id='late_capas_job',
                replace_existing=True
            )
            
            scheduler.add_job(
                check_and_send_reminders,
                'cron',
                hour=11,
                minute=59,
                id='reminders_morning_job',
                replace_existing=True
            )
            
            scheduler.add_job(
                check_and_send_reminders,
                'cron',
                hour=23,
                minute=59,
                id='reminders_evening_job',
                replace_existing=True
            )
            
            # Add a health check job
            scheduler.add_job(
                print_scheduler_jobs,
                'interval',
                hours=1,
                id='health_check_job',
                replace_existing=True
            )
            
            scheduler.start()
            print_scheduler_jobs()
            logging.info("Scheduler started successfully")
            
        except Exception as e:
            logging.error(f"Failed to start scheduler: {str(e)}", exc_info=True)
            raise
        
        atexit.register(lambda: scheduler.shutdown(wait=False))

# In your Flask app
with app.app_context():
    if not is_running_from_reloader():
        init_scheduler()





@app.errorhandler(404)
def not_found_error(error):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    return render_template('500.html'), 500

@app.errorhandler(403)
def forbidden_error(error):
    return render_template('403.html'), 403
if __name__ == '__main__':
    app.run(debug=False)


